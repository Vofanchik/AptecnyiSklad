import datetime

from openpyxl import load_workbook
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Mm


class Word_export:
    def form_docx(self, inf_to_fill, file_name):
        document = Document()

        section = document.sections[-1]
        section.page_width = Mm(210)
        section.page_height = Mm(297)

        p = document.add_paragraph('Остатки по состоянию на {}'.format(datetime.date.today()))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table = document.add_table(rows=1, cols=4, style='Table Grid')

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '№ п/п'
        hdr_cells[1].text = 'Наименование'
        hdr_cells[2].text = 'Ед. измерения'
        hdr_cells[3].text = 'Кол-во'



        for pp, lst in enumerate(inf_to_fill):
            row_cells = table.add_row().cells
            row_cells[0].text = str(pp+1)
            row_cells[1].text = lst[0]
            row_cells[2].text = lst[1]
            row_cells[3].text = lst[3]


        for i, a in zip(range(0,3),[0.5, 10, 2, 1]):
            for cell in table.columns[i].cells:
                cell.width = Cm(a)

        document.save(file_name)

class XlsxImport:
    def __init__(self, file_name):
        self.file_name = file_name


    def import_into_list(self):
        wb2 = load_workbook(self.file_name)
        sheet = wb2.active
        rows = sheet.max_row
        # cols = sheet.max_column
        items = []
        for i in range(1, rows + 1):
            items_vals = []

            for j in range(2,5):
                items_vals.append(sheet.cell(row=i, column=j).value)
            items.append(items_vals)

        return items

# p=Word_export()
# p.form_docx([['ф',"б","с"],['ф',"б","с"],['ф',"б","с"]])
